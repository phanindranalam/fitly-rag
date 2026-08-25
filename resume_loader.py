"""Turn an uploaded resume into text, then into structured fields.

WHY PARSING IS ITS OWN MODULE AND NOT THREE LINES IN THE UI
-----------------------------------------------------------
Everything downstream is built on this text. If extraction silently returns
half a resume -- which is the normal failure mode for a two-column PDF, and
resume templates love two columns -- retrieval still runs, generation still
runs, and the app confidently tells someone they have no cloud experience
because the skills column landed in an image or got interleaved into
gibberish. The failure is invisible at every later stage. So parsing gets a
module, a quality check, and a warning the user actually sees.

WHY TWO PARSERS
---------------
LlamaParse is a layout-aware document parser: it understands columns,
tables and reading order, which is precisely what pypdf does not. pypdf
extracts text in the order the PDF's content stream happens to emit it,
which for a two-column layout can interleave the columns line by line.

But LlamaParse is a network call with an API key and a queue, and it can be
slow or down. A resume parser that fails when a third-party service is
having a bad afternoon is not a resume parser. So: LlamaParse when a key is
present and it works, pypdf/python-docx otherwise, and the app always tells
you which one ran. Degrade, never break.

PDF AND DOCX ONLY, ON PURPOSE
-----------------------------
.doc (the pre-2007 binary format) needs LibreOffice to read reliably, and
.txt/.rtf resumes are rare enough that supporting them adds surface area for
no real coverage. Two formats, both parsed well, beats five parsed badly.

WHY EXTRACTION IS DETERMINISTIC AND NOT AN LLM CALL
---------------------------------------------------
Skills, titles and years come out of a keyword taxonomy and regexes, not a
model. An LLM asked to "extract the skills from this resume" will produce a
reasonable-looking list that includes skills the person does not have,
because it is completing a pattern, not reading. For the one artifact the
whole app trusts, a boring exact substring match is the right tool: if
"Kubernetes" appears in the output, it appeared in the resume. The taxonomy
is the same 560-term one Fitly uses, so a nurse's resume yields nursing
terms and not a shrug.
"""

from __future__ import annotations

import os
import re
from dataclasses import dataclass, field

import vendor_matching as matching

SUPPORTED_EXTENSIONS = (".pdf", ".docx")

# Below this, the "text" we got back is a filename and some page furniture,
# not a resume. Real resumes run 2,000-8,000 characters.
MIN_USABLE_CHARS = 400


@dataclass
class ParsedResume:
    text: str
    parser: str = ""                       # which backend actually ran
    warnings: list[str] = field(default_factory=list)
    skills: list[str] = field(default_factory=list)
    competencies: dict[str, list[str]] = field(default_factory=dict)
    titles: list[str] = field(default_factory=list)
    years_experience: float | None = None
    seniority: str = ""

    @property
    def usable(self) -> bool:
        return len(self.text) >= MIN_USABLE_CHARS

    def summary(self) -> str:
        bits = [f"{len(self.text):,} chars via {self.parser}",
                f"{len(self.skills)} skill term(s)"]
        if self.years_experience is not None:
            bits.append(f"~{self.years_experience:g} yrs")
        if self.seniority:
            bits.append(self.seniority)
        return ", ".join(bits)


# ---------------------------------------------------------------------------
# Text extraction
# ---------------------------------------------------------------------------

def _parse_with_llamaparse(path: str) -> str:
    """Layout-aware parse. Raises on any problem so the caller falls back."""
    key = os.getenv("LLAMA_CLOUD_API_KEY", "").strip()
    if not key:
        raise RuntimeError("no LLAMA_CLOUD_API_KEY")

    from llama_parse import LlamaParse

    parser = LlamaParse(
        api_key=key,
        # Markdown keeps headings and bullet structure, which the section
        # splitter downstream uses to tell "Experience" from "Education".
        # Plain text would flatten exactly the structure we paid for.
        result_type="markdown",
        # "fast" skips the multimodal path. Resumes are text documents; the
        # expensive mode is for scanned pages and complex tables, and 40k
        # credits go a lot further in fast mode.
        parse_mode="parse_page_without_llm",
        num_workers=1,
        verbose=False,
    )
    docs = parser.load_data(path)
    return "\n\n".join(d.text for d in docs).strip()


def _parse_pdf_local(path: str) -> str:
    from pypdf import PdfReader

    reader = PdfReader(path)
    return "\n\n".join((page.extract_text() or "") for page in reader.pages).strip()


def _parse_docx_local(path: str) -> str:
    import docx

    d = docx.Document(path)
    parts = [p.text for p in d.paragraphs]
    # Resume templates put contact details and skills grids in tables, and a
    # paragraph-only walk misses every one of them.
    for table in d.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(p for p in parts if p.strip()).strip()


_WS = re.compile(r"[ \t\xa0]+")
_NL = re.compile(r"\n{3,}")


def _clean(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = _WS.sub(" ", text)
    text = "\n".join(line.strip() for line in text.split("\n"))
    return _NL.sub("\n\n", text).strip()


def extract_text(path: str, prefer_llamaparse: bool = True) -> tuple[str, str, list[str]]:
    """Returns (text, parser_name, warnings)."""
    ext = os.path.splitext(path)[1].lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Unsupported file type {ext!r}. Upload a PDF or DOCX.")

    warnings: list[str] = []

    if prefer_llamaparse:
        try:
            text = _clean(_parse_with_llamaparse(path))
            if len(text) >= MIN_USABLE_CHARS:
                return text, "LlamaParse (layout-aware)", warnings
            warnings.append("LlamaParse returned very little text; fell back to local parsing.")
        except Exception as exc:
            # Not surfaced as an error: the fallback is expected to work, and
            # a user does not care which of two parsers ran.
            warnings.append(f"LlamaParse unavailable ({type(exc).__name__}); used the local parser.")

    text = _clean(_parse_pdf_local(path) if ext == ".pdf" else _parse_docx_local(path))
    return text, "pypdf" if ext == ".pdf" else "python-docx", warnings


# ---------------------------------------------------------------------------
# Structure
# ---------------------------------------------------------------------------

# Ordered longest-first so "senior staff engineer" wins over "engineer".
_TITLE_WORDS = [
    "chief", "vp", "vice president", "head of", "director", "principal",
    "staff", "senior", "sr.", "sr", "lead", "manager", "supervisor",
    "architect", "engineer", "developer", "analyst", "scientist",
    "consultant", "specialist", "administrator", "coordinator", "associate",
    "nurse", "practitioner", "technician", "designer", "researcher",
    "accountant", "controller", "recruiter", "teacher", "counselor",
    "paralegal", "attorney", "representative", "intern",
]
_TITLE_RE = re.compile(
    r"^\s*([A-Z][A-Za-z0-9&/,\.\-\+ ]{2,60}?)\s*(?:[|@,–—-]|$)", re.MULTILINE)

_YEARS_RE = re.compile(
    r"(\d{1,2})\+?\s*(?:\+)?\s*years?\s+(?:of\s+)?(?:progressive\s+|relevant\s+|professional\s+)?experience",
    re.IGNORECASE)

# 2019 - 2023, 2019 to Present, 03/2019 - 05/2023
_RANGE_RE = re.compile(
    r"(?:(\d{1,2})[/\-])?(\d{4})\s*(?:[–—\-]|to)\s*(?:(\d{1,2})[/\-])?(\d{4}|present|current|now)",
    re.IGNORECASE)

_SENIORITY = [
    ("executive", ("chief", "vp ", "vice president", "head of", "svp", "cto", "cio", "ceo")),
    ("director", ("director", "sr. director", "senior director")),
    ("manager", ("manager", "supervisor", "team lead")),
    ("senior", ("senior", "sr.", "staff", "principal", "lead ")),
    ("mid", ("engineer", "analyst", "specialist", "developer", "consultant")),
    ("entry", ("intern", "junior", "jr.", "associate", "trainee", "entry")),
]


def extract_titles(text: str, limit: int = 6) -> list[str]:
    """Lines that look like job titles.

    Heuristic on purpose. A line is a title if it starts a line, is short,
    and contains a known role word. That misses creative titles ("Growth
    Ninja") and it is supposed to: a wrong title poisons the search query,
    while a missing one just means the user types it themselves.
    """
    out, seen = [], set()
    for m in _TITLE_RE.finditer(text):
        cand = m.group(1).strip(" -|,")
        low = cand.lower()
        if len(cand) < 4 or len(cand) > 60:
            continue
        if not any(w in low for w in _TITLE_WORDS):
            continue
        if low in seen:
            continue
        seen.add(low)
        out.append(cand)
        if len(out) >= limit:
            break
    return out


def extract_years(text: str, current_year: int = 2026) -> float | None:
    """Years of experience, stated or inferred.

    A stated "12 years of experience" is taken at face value. Otherwise the
    span from the earliest to the latest employment year is used, which
    overcounts gaps and undercounts overlapping roles; it is reported as
    approximate everywhere it is shown for exactly that reason.
    """
    stated = [int(m.group(1)) for m in _YEARS_RE.finditer(text)]
    if stated:
        return float(max(stated))

    years: list[int] = []
    for m in _RANGE_RE.finditer(text):
        start = int(m.group(2))
        end_raw = m.group(4).lower()
        end = current_year if end_raw in ("present", "current", "now") else int(end_raw)
        # 1960 is a graduation year or a typo, not a career start.
        if 1980 <= start <= current_year and start <= end <= current_year:
            years.extend([start, end])
    if not years:
        return None
    span = max(years) - min(years)
    return float(span) if 0 < span <= 50 else None


def infer_seniority(titles: list[str], years: float | None) -> str:
    blob = " ".join(titles).lower()
    for label, markers in _SENIORITY:
        if any(mk in blob for mk in markers):
            return label
    if years is None:
        return ""
    if years >= 12:
        return "senior"
    if years >= 5:
        return "mid"
    return "entry"


# Acronym matching is what lets the taxonomy catch "SOC 2", "HIPAA" and
# "GAAP" out of running text, and the price is that a two-letter state code
# in an address line collides with a real term ("GA" the state vs "GA" the
# analytics tool). Dropping the colliding codes loses a genuine skill only
# for someone who listed the tool by acronym and never once by name, which
# is rarer than having an address.
_STATE_CODE_COLLISIONS = {
    "ga", "or", "in", "me", "hi", "id", "la", "ma", "md", "co", "de", "pa",
    "va", "wa", "ca", "ok", "ne", "mo", "mt", "nd", "sd",
}


def structure(text: str) -> dict:
    """Skills, titles, years, seniority -- all from the text, none invented."""
    skills = sorted(s for s in matching.extract_keywords(text)
                    if s not in _STATE_CODE_COLLISIONS)
    groups: dict[str, list[str]] = {}
    for group, terms in matching.COMPETENCY_GROUPS.items():
        hit = sorted(set(terms) & set(skills))
        if hit:
            groups[group] = hit
    titles = extract_titles(text)
    years = extract_years(text)
    return {
        "skills": skills,
        "competencies": groups,
        "titles": titles,
        "years_experience": years,
        "seniority": infer_seniority(titles, years),
    }


def load(path: str, prefer_llamaparse: bool = True) -> ParsedResume:
    text, parser, warnings = extract_text(path, prefer_llamaparse)
    parsed = ParsedResume(text=text, parser=parser, warnings=list(warnings))

    if not parsed.usable:
        parsed.warnings.append(
            f"Only {len(text)} characters came out of this file. If it is a scanned "
            "image or a heavily designed template, export a plain PDF from Word or "
            "Google Docs and try again -- otherwise the match will be wrong.")
        return parsed

    fields = structure(text)
    parsed.skills = fields["skills"]
    parsed.competencies = fields["competencies"]
    parsed.titles = fields["titles"]
    parsed.years_experience = fields["years_experience"]
    parsed.seniority = fields["seniority"]

    # A resume that yields almost no recognized terms usually means broken
    # extraction, not a candidate with no skills. Say so rather than quietly
    # producing a bad match -- this is the two-column-PDF failure surfacing.
    if len(parsed.skills) < 5:
        parsed.warnings.append(
            f"Only {len(parsed.skills)} recognizable skill term(s) were found. The file "
            "may have parsed badly (columns, text-as-image), which would make any "
            "match unreliable.")
    return parsed


def to_query(parsed: ParsedResume, target_role: str = "", top_skills: int = 12) -> str:
    """Build the retrieval query from a resume.

    A whole resume is a terrible query: 5,000 characters of education,
    addresses and hobbies embedded into one 384-dimension vector is a blurry
    average of a person, and it retrieves blurry averages of job postings.
    The query is the strongest signal only -- the target role if the user
    gave one, plus the highest-value skill terms -- because a query should
    look like the thing you want to find.
    """
    parts = []
    if target_role.strip():
        parts.append(target_role.strip())
    elif parsed.titles:
        parts.append(parsed.titles[0])
    # Multi-word terms first: "machine learning" discriminates between
    # postings far better than "python", which appears everywhere.
    ranked = sorted(parsed.skills, key=lambda s: (-len(s.split()), -len(s)))
    parts.extend(ranked[:top_skills])
    return ", ".join(parts)


if __name__ == "__main__":
    import json
    import sys

    if len(sys.argv) < 2:
        raise SystemExit("usage: python resume_loader.py <resume.pdf|resume.docx>")
    r = load(sys.argv[1])
    print(r.summary())
    for w in r.warnings:
        print(f"  ! {w}")
    print(json.dumps({
        "titles": r.titles,
        "years_experience": r.years_experience,
        "seniority": r.seniority,
        "competencies": r.competencies,
    }, indent=2))
    print("\nquery ->", to_query(r))
