"""Generate synthetic resumes for testing the matching path.

WHY SYNTHETIC
-------------
Testing the resume path with your own resume proves one thing: it works on
your resume. It cannot tell you what happens when the parser meets a
two-column layout, or when someone from outside the corpus's industry uploads
a document, or whether a junior and a staff engineer actually get different
answers. Those are the cases that break in front of an audience.

Five fictional people, chosen so each one probes a different failure:

  1  platform_sre.pdf      Dense overlap with the corpus. The happy path, and
                           the one worth demoing.
  2  data_ml.docx          Different domain, same industry. Also tests DOCX,
                           which takes a completely separate parser branch.
  3  junior_frontend.pdf   1.5 years. Should surface gaps against senior
                           postings rather than inventing a match.
  4  nurse.pdf             The taxonomy has 58 clinical terms, so skills
                           extract cleanly -- and the corpus has no nursing
                           jobs at all. The correct behaviour is an honest
                           "nothing here fits", not a stretched match. This is
                           the most valuable file in the set.
  5  two_column.pdf        Deliberately hostile layout. pypdf interleaves the
                           columns and the text comes out scrambled. The app
                           should say which parser ran and warn, because a
                           badly-parsed resume produces a confident wrong
                           match, which is worse than a failure.

Every person, employer, phone number and address here is invented. Emails use
example.com, phone numbers use the 555 reserved range.

    python make_test_resumes.py            # writes data/test_resumes/
"""

from __future__ import annotations

import os

OUT_DIR = os.path.join("data", "test_resumes")

FOOTER = "Synthetic test document. Not a real person."


# ---------------------------------------------------------------------------
# Content
# ---------------------------------------------------------------------------

PLATFORM_SRE = {
    "name": "Priya Raghunathan",
    "contact": "priya.raghunathan@example.com  |  (555) 0142  |  Atlanta, GA",
    "headline": "Staff Site Reliability Engineer",
    "sections": [
        ("SUMMARY", [
            "Staff SRE with 11 years running production infrastructure for high-traffic",
            "consumer platforms. Led reliability for a fleet of 160+ Kubernetes clusters",
            "across three regions. Comfortable owning on-call, incident command, and the",
            "unglamorous work of making postmortems change something.",
        ]),
        ("EXPERIENCE", [
            "Staff Site Reliability Engineer, Northwind Logistics  2021 - present",
            "  - Operate Kubernetes at scale: 160+ clusters, multi-region, zero-downtime upgrades.",
            "  - Rebuilt the on-call rotation and cut pages per engineer per week from 14 to 3.",
            "  - Own Terraform modules used by 40 engineers; migrated 300+ resources to modules.",
            "  - Built SLO tooling on Prometheus and Grafana; error budgets now gate releases.",
            "",
            "Senior Infrastructure Engineer, Calder Systems  2017 - 2021",
            "  - Ran the AWS estate: EKS, RDS, S3, IAM. Reduced spend 31% without cutting capacity.",
            "  - Wrote the incident response runbook and ran incident command for Sev1s.",
            "  - Introduced Docker-based CI, cutting build times from 22 minutes to 6.",
            "",
            "Infrastructure Engineer, Hartwell Payments  2015 - 2017",
            "  - Linux systems administration, Ansible configuration management, Nagios monitoring.",
        ]),
        ("SKILLS", [
            "Kubernetes, Docker, Terraform, AWS, Linux, Go, Python, Bash",
            "Prometheus, Grafana, Datadog, PagerDuty, ArgoCD, Helm",
            "CI/CD, GitOps, incident response, on-call, capacity planning, SRE practices",
            "PostgreSQL, Redis, Kafka",
        ]),
        ("EDUCATION", [
            "B.S. Computer Engineering, Georgia Institute of Technology, 2015",
        ]),
    ],
}

DATA_ML = {
    "name": "Marcus Deleon",
    "contact": "marcus.deleon@example.com  |  (555) 0177  |  Remote (US)",
    "headline": "Senior Data Engineer",
    "sections": [
        ("SUMMARY", [
            "Data engineer, 7 years, focused on batch and streaming pipelines feeding",
            "machine learning systems. Most recently owned the feature store that four",
            "model teams depend on.",
        ]),
        ("EXPERIENCE", [
            "Senior Data Engineer, Bellrose Analytics  2020 - present",
            "  - Built and operate 90+ Airflow DAGs moving 4TB/day into Snowflake.",
            "  - Designed the feature store used by four ML teams; cut training data prep from days to hours.",
            "  - Streaming ingestion with Kafka and Spark Structured Streaming.",
            "  - dbt for transformations; introduced testing and cut silent data quality failures.",
            "",
            "Data Engineer, Kestrel Retail Group  2018 - 2020",
            "  - ETL in Python and SQL against PostgreSQL and Redshift.",
            "  - Partnered with data scientists on scikit-learn and PyTorch model deployment.",
        ]),
        ("SKILLS", [
            "Python, SQL, Scala, Spark, Airflow, dbt, Kafka",
            "Snowflake, Redshift, PostgreSQL, BigQuery",
            "PyTorch, scikit-learn, pandas, numpy, MLflow",
            "AWS, Docker, Terraform (working knowledge)",
        ]),
        ("EDUCATION", [
            "M.S. Statistics, University of Illinois, 2018",
            "B.S. Mathematics, University of Illinois, 2016",
        ]),
    ],
}

JUNIOR_FRONTEND = {
    "name": "Tobi Aderinto",
    "contact": "tobi.aderinto@example.com  |  (555) 0198  |  Chicago, IL",
    "headline": "Frontend Developer",
    "sections": [
        ("SUMMARY", [
            "Frontend developer with 1.5 years of professional experience building",
            "React interfaces. Bootcamp graduate, previously in retail management.",
        ]),
        ("EXPERIENCE", [
            "Frontend Developer, Marlow Interactive  2025 - present",
            "  - Build and maintain React and TypeScript components for a customer dashboard.",
            "  - Improved Lighthouse performance score from 61 to 88 on the main dashboard.",
            "  - Write unit tests in Jest; participate in code review.",
            "",
            "Junior Developer (contract), Fennimore Studio  2024 - 2025",
            "  - HTML, CSS, JavaScript work on marketing sites. Some Vue.",
        ]),
        ("SKILLS", [
            "JavaScript, TypeScript, React, HTML, CSS, Tailwind",
            "Git, Jest, Figma, Vite",
            "Agile, code review",
        ]),
        ("EDUCATION", [
            "Full-Stack Web Development Certificate, Ravenwood Bootcamp, 2024",
            "B.A. Communications, DePaul University, 2019",
        ]),
    ],
}

NURSE = {
    "name": "Denise Okafor, RN, BSN",
    "contact": "denise.okafor@example.com  |  (555) 0163  |  Columbus, OH",
    "headline": "Registered Nurse - Critical Care",
    "sections": [
        ("SUMMARY", [
            "Registered nurse with 9 years in critical care. ICU and emergency",
            "experience, charge nurse for a 24-bed unit, preceptor for new graduates.",
        ]),
        ("EXPERIENCE", [
            "Charge Nurse, ICU, Cairnstone Regional Medical Center  2019 - present",
            "  - Direct patient care and patient assessment for a 24-bed intensive care unit.",
            "  - Triage, acute care, ventilator management, medication administration.",
            "  - Precept new graduate nurses; lead shift handoff and care coordination.",
            "  - EPIC documentation; participate in quality improvement and infection control.",
            "",
            "Staff Nurse, Emergency Department, Waverly Memorial  2016 - 2019",
            "  - Emergency triage, trauma response, IV therapy, wound care.",
        ]),
        ("SKILLS", [
            "Patient care, patient assessment, triage, ICU, emergency, acute care",
            "Medication administration, IV therapy, wound care, care coordination",
            "EPIC, BLS, ACLS, PALS, infection control, quality improvement",
        ]),
        ("EDUCATION", [
            "B.S.N. Nursing, Ohio State University, 2016",
        ]),
    ],
}

TWO_COL_LEFT = [
    "Hal Brantley",
    "DevOps Engineer",
    "",
    "CONTACT",
    "hal.brantley@example.com",
    "(555) 0129",
    "Denver, CO",
    "",
    "SKILLS",
    "Kubernetes",
    "Docker",
    "Jenkins",
    "Terraform",
    "AWS",
    "Python",
    "Bash",
    "Ansible",
    "GitLab CI",
    "Nagios",
    "",
    "EDUCATION",
    "B.S. Information",
    "Systems",
    "Colorado State, 2017",
]

TWO_COL_RIGHT = [
    "EXPERIENCE",
    "",
    "DevOps Engineer",
    "Ashgrove Media  2021 - present",
    "- Maintain CI/CD pipelines in GitLab CI for 30 services.",
    "- Manage EKS clusters and Terraform state for three environments.",
    "- Reduced deploy failures 40% by adding smoke tests to the pipeline.",
    "- Rotating on-call, one week in four.",
    "",
    "Systems Engineer",
    "Pinehill Data  2018 - 2021",
    "- Linux server administration and Ansible playbooks.",
    "- Migrated 60 VMs from on-premise VMware to AWS EC2.",
    "- Built monitoring with Nagios and later Prometheus.",
    "",
    "IT Support Specialist",
    "Colorado State University  2017 - 2018",
    "- Desktop support, Active Directory, ticket triage.",
]


# ---------------------------------------------------------------------------
# Writers
# ---------------------------------------------------------------------------

def write_pdf(spec: dict, path: str) -> None:
    from reportlab.lib.pagesizes import LETTER
    from reportlab.pdfgen import canvas

    c = canvas.Canvas(path, pagesize=LETTER)
    width, height = LETTER
    x, y = 54, height - 60

    c.setFont("Helvetica-Bold", 17)
    c.drawString(x, y, spec["name"])
    y -= 19
    c.setFont("Helvetica", 11)
    c.drawString(x, y, spec["headline"])
    y -= 15
    c.setFont("Helvetica", 9)
    c.drawString(x, y, spec["contact"])
    y -= 26

    for title, lines in spec["sections"]:
        if y < 90:
            c.showPage()
            y = height - 60
        c.setFont("Helvetica-Bold", 10.5)
        c.drawString(x, y, title)
        y -= 5
        c.line(x, y, width - 54, y)
        y -= 14
        c.setFont("Helvetica", 9.5)
        for line in lines:
            if y < 70:
                c.showPage()
                y = height - 60
                c.setFont("Helvetica", 9.5)
            c.drawString(x, y, line)
            y -= 12.5
        y -= 10

    c.setFont("Helvetica-Oblique", 7.5)
    c.drawString(x, 40, FOOTER)
    c.save()


def write_docx(spec: dict, path: str) -> None:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    h = doc.add_paragraph()
    run = h.add_run(spec["name"])
    run.bold = True
    run.font.size = Pt(17)

    doc.add_paragraph(spec["headline"])
    doc.add_paragraph(spec["contact"])

    for title, lines in spec["sections"]:
        p = doc.add_paragraph()
        r = p.add_run(title)
        r.bold = True
        r.font.size = Pt(11)
        for line in lines:
            doc.add_paragraph(line)

    doc.add_paragraph(FOOTER)
    doc.save(path)


def write_two_column_pdf(path: str) -> None:
    """A layout designed to be hard to parse.

    pypdf extracts by text-object order, not by visual column, so a two-column
    resume comes out interleaved: a skill, then a bullet from the other side,
    then another skill. The extracted string is nearly unreadable and any
    skill list built from it is unreliable. That is the point -- the app must
    notice and say so rather than confidently matching on scrambled text.
    """
    from reportlab.lib.pagesizes import LETTER
    from reportlab.pdfgen import canvas

    c = canvas.Canvas(path, pagesize=LETTER)
    width, height = LETTER
    top = height - 60

    # Deliberately interleave the two columns' draw order, one line at a time.
    left_y = top
    right_y = top
    for i in range(max(len(TWO_COL_LEFT), len(TWO_COL_RIGHT))):
        if i < len(TWO_COL_LEFT):
            line = TWO_COL_LEFT[i]
            bold = line.isupper() and len(line) > 2
            c.setFont("Helvetica-Bold" if bold else "Helvetica", 9)
            c.drawString(50, left_y, line)
            left_y -= 13
        if i < len(TWO_COL_RIGHT):
            line = TWO_COL_RIGHT[i]
            bold = line.isupper() and len(line) > 2
            c.setFont("Helvetica-Bold" if bold else "Helvetica", 9)
            c.drawString(235, right_y, line)
            right_y -= 13

    c.setFont("Helvetica-Oblique", 7.5)
    c.drawString(50, 40, FOOTER)
    c.save()


# ---------------------------------------------------------------------------

FILES = [
    ("platform_sre.pdf",    lambda p: write_pdf(PLATFORM_SRE, p)),
    ("data_ml.docx",        lambda p: write_docx(DATA_ML, p)),
    ("junior_frontend.pdf", lambda p: write_pdf(JUNIOR_FRONTEND, p)),
    ("nurse.pdf",           lambda p: write_pdf(NURSE, p)),
    ("two_column.pdf",      write_two_column_pdf),
]


def main() -> None:
    os.makedirs(OUT_DIR, exist_ok=True)
    for name, fn in FILES:
        path = os.path.join(OUT_DIR, name)
        fn(path)
        print(f"{os.path.getsize(path):7,d} bytes  {path}")
    print(f"\n{len(FILES)} synthetic resumes in {OUT_DIR}")
    print("Test them:  python ui_test.py --resumes " + OUT_DIR)


if __name__ == "__main__":
    main()
